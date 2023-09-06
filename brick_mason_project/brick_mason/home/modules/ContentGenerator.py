#from pprint import pprint
#import nltk
#from Questgen import main
import tkinter as tk
from tkinter import filedialog as fd
from tkinter.messagebox import showinfo
import ctypes
import os
import docx
#user32 = ctypes.windll.user32
from home.modules.SourceImporter import extract_from_docx
import argparse
import os
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lex_rank import LexRankSummarizer

#

def summarize_file(f_record, sentences_count):
    if f_record.file.name.endswith('.txt'):
        with open(f_record.file, "r") as file: #encoding="iso-8859-1"
            rtnd_text = file.read()
    elif f_record.file.name.endswith('.docx'):
        rtnd_text = extract_from_docx(f_record.file)
    else:
        raise ValueError("Unsupported file type: {}".format(f_record.file.name))
    parser = PlaintextParser.from_string(rtnd_text, Tokenizer("english"))
    summarizer = LexRankSummarizer()
    summary = summarizer(parser.document, sentences_count=sentences_count)
    return " ".join([str(sentence) for sentence in summary])

''' Remove Before Demo, uncomment QuestGen


def payload(text): return {"input_text": text}

# Generates Boolean question and answers from provided text
# Returns list [].
def generate_bq(text):
    qg = main.BoolQGen()
    text_in = payload(text)
    output = qg.predict_boolq(text_in)
    return output['Boolean Questions']

#
def generate_mcq(text):
    qg = main.QGen()
    text_in = payload(text)
    output = qg.predict_mcq(text_in)
    return output


qe = main.BoolQGen()

doc = docx.Document('chemComputing.docx')

# Add all paragraphs to a list
full_text = []
for para in doc.paragraphs:
    full_text.append(para.text)

# Handle the text -- REPLACE: prob need to loop through paragraphs and analyze each one individually as opposed to join. Char limit for model is 512
text = {
    "input_text": " ".join(full_text)
}

# Q and A code generation code obtained at : https://github.com/ramsrigouthamg/Questgen.ai - begin quote

# Boolean questions - potential to use as an objective list and as titles in the formatted draft
output = qe.predict_boolq(text)
pprint(output)

# Multiple choice questions - append to end of draft
qg = main.QGen()
output = qg.predict_mcq(text)
pprint(output)

# Generate FAQs - also potential for section titles or objectives in the formatted draft
output = qg.predict_shortq(text)
pprint(output)

# Incorporate if possible - allows user to request rewording of a questions
oldPhrase = {
    "input_text": "What are disadvantages of traditional computers?",
    "max_questions": 5
}
output = qg.paraphrase(oldPhrase)
pprint(output)

'''






""" docx-bookmark: docx bookmarks
Functions addFlexTable, addPlot , addParagraph and addImage can send respective outputs into these bookmarks.
When used with addPlot, addParagraph and addImage, content (plots, paragraphs or images) will replace the whole paragraph containing the bookmark.
When used with addFlexTable content (table) will be inserted after the paragraph containing the bookmark.
"""

"""
def bookmark_funct():
    wordApp = user32.gencache.EnsureDispatch(
    'Word.Application')  # create a word application object
    wordApp.Visible = False  # hide the word application

    doc = wordApp.Documents.Open("template.doc")  # opening the template file

    # for creating a new one: doc = wordApp.Documents.Add()

    # change the string Name to whatever name of your bookmarks
    rng=doc.Bookmarks("Name").Range
    rng.InsertAfter("Hello World")
    rng=doc.Bookmarks("Picture").Range
    rng.InlineShapes.AddPicture("/path/to/your/picture")
    doc.SaveAs("new_file.doc")
    
    #----------------

    #change bookmark text, replacing bookmark
    Word.Document doc = new Word.Document();
    Word.Application app = new Word.Application();
    app.Visible = true;
    doc = app.Documents.Open("D:\\test.docx");
    Word.Range rng = doc.Bookmarks[1].Range;
    rng.Text = "Hello ";
    doc.Bookmarks.Add("BookMark2", rng); 


# SAMPLE GET BOOKMARKS FUNCTION
# req: install docx
# document is docx.Document('path_example.docx')
def get_bookmarks(document):
    doc_element = document.part._element
    bookmarks_list = doc_element.findall('.//' + qn('w:bookmarkStart'))
    bookmarks_text = []

    for bookmark in bookmarks_list:
        par = bookmark.getparent()
        runs = par.findall(qn('w:r'))
        for run in runs:
            try:
                bookmarks_text.append(run.find(qn('w:t')).text)
            except Exception as e:
                print(e)

    return bookmarks_text
"""
