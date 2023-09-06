
from home.models import Brick, ReferenceList, TermList
import docx
import tkinter as tk
from tkinter import simpledialog, filedialog, messagebox
from home.modules.ContentGenerator import summarize_file
from django.core.files import File
from brick_mason.settings import S3_R, AWS_STORAGE_BUCKET_NAME
import botocore
import os
from pathlib import Path

'''====================================== General Functions ======================================='''

''' Searchs a queryset for the title, and returns local drive path of brick '''
def get_record_path(brick_list, title):
    brick = brick_list.get(title=title)
    path = brick.workspace + "/" + title + ".docx"
    return path

# ----------------------------------------------------------------------------------------------------------------

''' Takes SourceFile object, and returns sentence count of 10 sentences,
    or 1 sentence/100 words, whichever is highest. '''
def get_sentence_count(record):
    count = (record.word_count / 100)
    if (10 < count):
        return count
    else:
        return 10

# ----------------------------------------------------------------------------------------------------------------

def no_brick_loaded():
    root = tk.Tk()
    root.withdraw()
    root.attributes("-topmost", True)
    messagebox.showinfo(title='No Brick loaded',
                        message='No Brick has been loaded, please load Brick.')
    root.destroy()

'''========================================= File Dropdown =========================================='''

'''
def save_local_file(path):
    doc = docx.Document(path)
    parts = path.split('.')
    newpath = parts[0] + "|" + parts[1]
    doc.save(newpath)
    doc.close()
    os.remove(path)
    os.rename(newpath, path)
'''

# Saves current brick in db and local drive.
def brick_save(file):
    path = Path(file)
    name = 'bricks/' + path.name.replace(' ', '_')
    title = path.name.split(".")[0]
    
    # Replace db file with new file.
    try:
        brick = Brick.objects.get(title=title)
    
        with path.open(mode='rb') as f:
            brick.file = File(f, name=name)
            brick.save()
    except:
        print("Unable to save brick related to path ")
        print(path)

# ----------------------------------------------------------------------------------------------------------------
    
def brick_save_as(file):
    print("EMPTY brick_save_as STUB!!!")

# ----------------------------------------------------------------------------------------------------------------

def open_selected(query, i):
    global return_path
    return_path = None
    record = query[i]
    path = record.workspace + "/" + record.title + ".docx"
    try:
        open_file(path)
        return_path = path
    except:
        try:
            download_s3_file(record.file.name, path)
            open_file(path)
            return_path = path
        except:
            print("Unable to find brick in local drive or S3.")
            return_path = None

# ----------------------------------------------------------------------------------------------------------------

def select_brick_to_open(q_list, query):
    len_max = 0
    for m in q_list:
        if len(m) > len_max:
            len_max = len(m)
    
    root = tk.Tk()
    root.geometry("700x300")
    root.title("Select ")
    root.attributes("-topmost", True)

    list = tk.Listbox(root, width = len_max)
    list.insert(0, *q_list)

    def b_trig():
        open_selected(query, list.curselection()[0])
        root.destroy()

    # close window
    def c_trig():
        global return_path
        return_path = None
        root.destroy()

    b = tk.Button(root, text="Select", command=b_trig)
    c = tk.Button(root, text="Cancel", command=c_trig)

    list.pack()
    b.pack(side=tk.LEFT)
    c.pack(side=tk.RIGHT)
    root.mainloop()

# ----------------------------------------------------------------------------------------------------------------

def get_stage(stage):
    if stage == 1:
        return "Brick Drafting"
    elif stage == 2:
            return "SME Review"
    elif stage == 3:
        return "Editor Review"
    elif stage == 4:
        return "Ready For Insert"
    elif stage == 5:
        return "Complete"
    else:
        return "Err: Out Of Range"

# ----------------------------------------------------------------------------------------------------------------

def query_to_list(query):
    list = []
    for record in query:
        stage = get_stage(record.stage)
        list.append("Creator: " + record.uid + " | " +
                    "Title: " + record.title + " | " + "Stage: "
                     + stage + " | " + "Workspace: "
                     + record.workspace)
    return list
    
# ----------------------------------------------------------------------------------------------------------------

''' Option 3: Load. Creates a window with a list of Brick '''
def brick_load(path):
    
    # if file was opened, request save
    if not ((path == "") or (path is None) or (path == "None")):
        brick_save(path)
        
    # Replace db file with new file.
    bricks = Brick.objects.all()
    
    # Convert query list to list of strings for display.
    brick_list = query_to_list(bricks)
    
    # Create window for selection and loading of brick.
    select_brick_to_open(brick_list, bricks)
    
    if (return_path is None) and not (path is None): 
        return path
    else:
        return return_path

# ----------------------------------------------------------------------------------------------------------------
    
def import_brick(file):
    print("EMPTY import_brick STUB!!!")

# ----------------------------------------------------------------------------------------------------------------
            
def export_brick(file):
    print("EMPTY export_brick STUB!!!")

# ----------------------------------------------------------------------------------------------------------------

def rename_local_file(path, new_title):
    try:
        dir = os.path.dirname(path)
        new_path = dir + "/" + new_title + ".docx"
        os.rename(path, new_path)
        return Path(new_path)
    except:
        print("rename_local_file(): Unable to rename local file.")

# ----------------------------------------------------------------------------------------------------------------

# Renames brick file in both db and local drive.
def rename_brick(path):
    path = Path(path)
    #name = 'bricks/' + path.name.replace(' ', '_')
    old_title = path.name.split(".")[0]
    
    new_title = request_title()
    if new_title == "": return
    
    path = rename_local_file(path, new_title)
    
    # Replace db file with new file.
    try:
        brick = Brick.objects.get(title=old_title)
        with path.open(mode='rb') as f:
            brick.file = File(f, name='bricks/'+path.name)
            brick.title = new_title
            brick.save()
            return path
    except:
        print("rename_brick: rename_brick: Unable to rename brick related to path:")
        print(path)

# ----------------------------------------------------------------------------------------------------------------
        
def exit_process(file):
    print("EMPTY exit_process STUB!!!")

'''========================================= Insert Dropdown =========================================='''

def summarizer(path):
    orig_text = 'Summary point 1'
    
    brick_save(path)
    
    path = Path(path)
    doc = docx.Document(path)
    title = path.name.split(".")[0]

    # Replace db file with new file.
    try:
        brick = Brick.objects.get(title=title)
        summary_text = summarize_file(brick, 10)
        replace_brick_paragraph(doc, orig_text, summary_text)
        doc.save(path)
        brick_save(path)
    except:
        print("summarizer: Unable to summarize brick related to path:")
        print(path)

# ----------------------------------------------------------------------------------------------------------------

def qa_list(file):
    print("EMPTY qa_list STUB!!!")

# ----------------------------------------------------------------------------------------------------------------
    
def image(file):
    print("EMPTY image STUB!!!")

# ----------------------------------------------------------------------------------------------------------------
    
def source_link(file):
    print("EMPTY source_link STUB!!!")
    
'''========================================= View Dropdown =========================================='''


def references(file):
    print("EMPTY references STUB!!!")

# ----------------------------------------------------------------------------------------------------------------
    
def contradictions(file):
    print("EMPTY contradictions STUB!!!")

# ----------------------------------------------------------------------------------------------------------------
    
def redundancies(file):
    print("EMPTY redundancies STUB!!!")

# ----------------------------------------------------------------------------------------------------------------
    
def outdated(file):
    print("EMPTY outdated STUB!!!")

# ----------------------------------------------------------------------------------------------------------------
    
def side_by_side(file):
    print("EMPTY side_by_side STUB!!!")
    

# ----------------------------------------------------------------------------------------------------------------

def popout_reference(file):
    print("EMPTY popout_reference STUB!!!")
    

'''========================================= Tools Dropdown =========================================='''

def flag(path):
    print("EMPTY flag STUB!!!")

# ----------------------------------------------------------------------------------------------------------------

''' Saves selected brick stage. '''
def set_brick_stage(brick, i):
    brick.stage = i + 1
    brick.save()

# ----------------------------------------------------------------------------------------------------------------

''' Display pop-up with list of bricks to select, changes stage
    of currently open brick to user selected stage. '''
def select_brick_stage(q_list, brick):
    len_max = 0
    for m in q_list:
        if len(m) > len_max:
            len_max = len(m)

    root = tk.Tk()
    root.geometry("150x300")
    root.title("Select ")
    root.attributes("-topmost", True)

    list = tk.Listbox(root, width=len_max)
    list.insert(0, *q_list)

    def b_trig():
        set_brick_stage(brick, list.curselection()[0])
        root.destroy()

    # close window
    def c_trig():
        root.destroy()

    b = tk.Button(root, text="Select", command=b_trig)
    c = tk.Button(root, text="Cancel", command=c_trig)

    list.pack()
    b.pack(side=tk.LEFT)
    c.pack(side=tk.RIGHT)
    root.mainloop()

# ----------------------------------------------------------------------------------------------------------------


def set_stage(path):
    path = Path(path)
    title = path.name.split(".")[0]
    
    # Replace db file with new file.
    try:
        stages = ["Brick Drafting", "SME Review", "Editor Review", "Ready For Insert", "Complete"]
        brick = Brick.objects.get(title=title)
        select_brick_stage(stages, brick)
    except:
        print("set_stage: Unable to save new stage to Brick record.")

# ----------------------------------------------------------------------------------------------------------------

'''' Set selected brick's workspace to user provided directory. '''
def set_brick_workspace(query, i):
    record = query[i]
    
    workspace = request_workspace()
    if workspace == "": return
    
    record.workspace
    record.save()    

# ----------------------------------------------------------------------------------------------------------------

''' Display pop-up with list of bricks to select, changes workspace
    of selected brick to user selected workspace'''
def change_brick_workspace(q_list, query):
    len_max = 0
    for m in q_list:
        if len(m) > len_max:
            len_max = len(m)

    root = tk.Tk()
    root.geometry("700x300")
    root.title("Select ")
    root.attributes("-topmost", True)

    list = tk.Listbox(root, width=len_max)
    list.insert(0, *q_list)

    def b_trig():
        set_brick_workspace(query, list.curselection()[0])
        root.destroy()

    # close window
    def c_trig():
        root.destroy()

    b = tk.Button(root, text="Select", command=b_trig)
    c = tk.Button(root, text="Cancel", command=c_trig)

    list.pack()
    b.pack(side=tk.LEFT)
    c.pack(side=tk.RIGHT)
    root.mainloop()

# ----------------------------------------------------------------------------------------------------------------

''' Allows selection of Brick from all on record, user provides
    replacement workspace from that Brick. '''
def set_workspace():

    # Replace db file with new file.
    bricks = Brick.objects.all()

    # Convert query list to list of strings for display.
    brick_list = query_to_list(bricks)
    
    # Select Brick from list and change the workspace.
    try:
        change_brick_workspace(brick_list, bricks)
    except:
        print("set_workspace: Unable to save new workspace to Brick record.")
        

'''====================================== New Brick Generation ======================================='''

'''  Request workspace for brick from user. '''
def request_workspace():
    workspace = ""

    try:
        root = tk.Tk()
        root.withdraw()
        root.attributes("-topmost", True)
        workspace = filedialog.askdirectory(title="Select Workspace")
        root.destroy()        
    except:
        print("Error Occured in request_workspace().")  
        return ""
    
    if workspace is None:
        return ""
    
    return workspace

# ----------------------------------------------------------------------------------------------------------------

'''  Request title for brick from user. '''
def request_title():
    title = ""

    try:
        root = tk.Tk()
        root.withdraw()
        root.attributes("-topmost", True)
        title = simpledialog.askstring(title="Brick Title",
                                       prompt="What would you like to title your new Brick?")
        root.destroy()
    except:
        print("Error Occured in request_title().")
        return ""

    if title is None:
        return ""
    
    return title

# ----------------------------------------------------------------------------------------------------------------

''' Sets up filename and path, then used to call
    download_s3_file(filename, path). '''
def get_copy_brick_template(title, workspace):
    worddoc = ""
    
    filename = "static/Brick Template.docx"
    path = workspace + '/' + title + '.docx'

    download_s3_file(filename, path)
    
    return path

# ----------------------------------------------------------------------------------------------------------------

''' Download S3 file by filename to path location on local path directory. '''
def download_s3_file(filename, path):
    try:
        S3_R.Bucket(AWS_STORAGE_BUCKET_NAME).download_file(filename, path)
    except botocore.exceptions.ClientError as e:
        if e.response['Error']['Code'] == "404":
            print("The object does not exist.")
        else:
            raise

# ----------------------------------------------------------------------------------------------------------------

''' generates new brick record, copies brick template document,
    sets title, and workspace. '''
def generate_new_brick(username):
    root = tk.Tk()
    root.withdraw()
    root.attributes("-topmost", True)
    
    # request title for brick, error or cancel returns empty string.
    workspace = request_workspace()
    if workspace == "": return ""

    # request title for brick, error or cancel returns empty string.
    title = request_title()
    if title == "": return ""
    
    # get downloaded file 
    path = Path(get_copy_brick_template(title, workspace))

    # Create and add new Brick record, stage 1 = 'Brick Drafting'
    with path.open(mode='rb') as f:
        brick_record = Brick(file=File(f, name='bricks/'+path.name), uid=username, title=title,
                         stage=1, workspace=workspace)
        brick_record.save()
    
    return {"record": brick_record, "path": path}

# ----------------------------------------------------------------------------------------------------------------

''' Uses provided file path to open file. '''
def open_file(path):
    if path is not (None or ''):
        os.startfile(os.path.abspath(path))

# ----------------------------------------------------------------------------------------------------------------

''' Creates a copy of the reference file and returns the path.'''
def make_working_file(record, workspace):
    name = record.fid.file.name
    path = workspace + '/' + name
    download_s3_file(name, path)
    
    return Path(path)

# ----------------------------------------------------------------------------------------------------------------

''' For each record in the query set, creats a working copy,
    transfers query to ReferenceList '''
def link_files_to_brick(queryset, brick_record):
    for record in queryset:
        if record.fid.file.name.endswith('.docx'):
            path = make_working_file(record, brick_record.workspace)
            with path.open(mode='rb') as f:
                rl = ReferenceList(bid_id=brick_record.pk, fid=record.fid,
                                   file=File(f, name= brick_record.title + "/" + path.name))
                rl.save()
            os.remove(path)
        else:
            rl = ReferenceList(bid_id=brick_record.pk, fid=record.fid,
                               file=None)
            rl.save()

# ----------------------------------------------------------------------------------------------------------------

''' Takes the unassigned terms in the db and assigns them to the
    new brick. '''
def link_terms_to_brick(username, bid):
    terms = TermList.objects.filter(bid=None, username=username)
    
    if terms.exists():
        for rec in TermList.objects.all():
            # Delete all records not attached to a brick.
            rec.bid_id = bid
            rec.save()

# ----------------------------------------------------------------------------------------------------------------

def insert_brick_image(doc, target_text, image):
    for paragraph in doc.paragraphs:
        if paragraph.text == target_text:
            r = paragraph.add_run()
            pic_p = r.add_picture(image)
            return pic_p

# ----------------------------------------------------------------------------------------------------------------

''' Provide docx Document obect, paragraph pointer, and text to insert.
    Inserts paragraph after target paragraph, and returns the paragraph
    pointer. '''
def add_next_brick_paragraph(doc, para_p, text):
    new_para_p = doc.add_paragraph(text)._p
    para_p.addnext(new_para_p)
    return new_para_p


# ----------------------------------------------------------------------------------------------------------------

''' Provide docx Document obect, text to search for and replacement text.
    Swaps original text with new text and returns the paragraph pointer. '''
def replace_brick_paragraph(doc, org_text, new_text):    
    for paragraph in doc.paragraphs:
        if paragraph.text == org_text:
            paragraph.text = new_text
            return paragraph._p


# ----------------------------------------------------------------------------------------------------------------

''' Takes the files, extracts information from them and inserts it
    into the new brick document. '''
def build_brick(queryset, path, title):
    doc = docx.Document(path)
    target = ['Title of Brick','Narrative text','Header 3. Narrative text']
    para_p = None
    
    # replace title of brick
    replace_brick_paragraph(doc, target[0], title)
    
    #insert summary text and pictures 
    for f_record in queryset:
        if f_record.fid.file.name.endswith(('.png', '.jpg')):
            image = f_record.fid.file
            insert_brick_image(doc, target[2], image)             
        elif f_record.fid.file.name.endswith('.docx'):
            sentence_count = get_sentence_count(f_record.fid)
            if (para_p is None):
                summary_text = summarize_file(f_record.fid, sentence_count)
                para_p = replace_brick_paragraph(doc, target[1], summary_text)
            else:
                summary_text = summarize_file(f_record.fid, sentence_count)
                para_p = add_next_brick_paragraph(doc, para_p, summary_text)      

    doc.save(path) 

# ----------------------------------------------------------------------------------------------------------------

''' Takes search result query and creates a new brick, updating and
    linking tables in the process. '''
def new_brick_generation(queryset, username):
    
    # return: {"record": brick_record, "file": downloaded_file}
    brick = generate_new_brick(username)
    if brick == "": return ""
    
    # Link Query Set to Brick
    link_files_to_brick(queryset, brick['record'])
    
    # Link Terms to Brick
    link_terms_to_brick(username, brick['record'].pk)
    
    #Insert brick text and pictures
    build_brick(queryset, brick['path'], brick['record'].title)
    
    # save changes to Brick in database.
    brick_save(brick['path'])
    
    #opens Brick Template file
    open_file(brick['path'])
    
    return brick['path']
